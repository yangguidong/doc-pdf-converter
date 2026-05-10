"""
Word-PDF interconversion tool.
Usage:
    py doc_pdf_converter.py input.docx              # -> input.pdf
    py doc_pdf_converter.py input.pdf               # -> input.docx
    py doc_pdf_converter.py input.docx output.pdf   # explicit output
    py doc_pdf_converter.py input.docx -o out.pdf   # explicit output
    py doc_pdf_converter.py --gui                   # graphical file picker
"""

import argparse
import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pathlib import Path
from typing import Optional


# ── Backend availability cache ───────────────────────────────────────

_AVAILABLE = {
    "word": None,        # MS Word COM
    "wps": None,         # WPS Writer COM
    "libreoffice": None, # LibreOffice CLI
}

# Try multiple ProgIDs for WPS Writer
_WPS_PROGIDS = ["KWPS.Application", "WPS.Application", "WPSWriter.Application"]


def _check_available():
    """Lazy detection: only checks libbreoffice command. COM is tried on-demand."""

    if _AVAILABLE["libreoffice"] is None:
        import subprocess
        try:
            result = subprocess.run(
                ["soffice", "--version"],
                capture_output=True, text=True, timeout=5,
            )
            _AVAILABLE["libreoffice"] = result.returncode == 0
        except Exception:
            _AVAILABLE["libreoffice"] = False


# ── Conversion functions ──────────────────────────────────────────────

def word_to_pdf(input_path: str, output_path: str) -> None:
    """Convert .docx to PDF using best available method."""

    _check_available()

    # Priority: WPS → Word → LibreOffice → pure Python
    if _AVAILABLE["wps"] is not False:
        try:
            progid = _AVAILABLE["wps"] if _AVAILABLE["wps"] else "KWPS.Application"
            _com_export(input_path, output_path, progid)
            _AVAILABLE["wps"] = progid
            print(f"[OK] WPS COM ({progid}) export successful")
            return
        except Exception as e:
            _AVAILABLE["wps"] = False
            print(f"[warn] WPS COM failed ({e}), falling back...")

    if _AVAILABLE["word"] is not False:
        try:
            _com_export(input_path, output_path, "Word.Application")
            _AVAILABLE["word"] = True
            print("[OK] MS Word COM export successful")
            return
        except Exception as e:
            _AVAILABLE["word"] = False
            print(f"[warn] Word COM failed ({e}), falling back...")

    if _AVAILABLE["libreoffice"]:
        try:
            _word_to_pdf_libreoffice(input_path, output_path)
            print("[OK] LibreOffice export successful")
            return
        except Exception as e:
            print(f"[warn] LibreOffice failed ({e}), falling back...")
            _AVAILABLE["libreoffice"] = False

    print("[info] Using pure Python fallback...")
    _word_to_pdf_pure(input_path, output_path)


def _com_export(input_path: str, output_path: str, progid: str) -> None:
    """Export docx to PDF via COM in a subprocess (isolates COM from main process)."""
    import subprocess
    import json
    import tempfile

    script = r"""
import sys, os, json
params = json.loads(sys.argv[1])
import pythoncom
pythoncom.CoInitialize()
try:
    import win32com.client
    app = win32com.client.Dispatch(params['progid'])
    app.Visible = False
    app.DisplayAlerts = 0
    doc = app.Documents.Open(params['input'], ReadOnly=True)
    try:
        doc.ExportAsFixedFormat(params['output'], ExportFormat=17,
                                OpenAfterExport=False, OptimizeFor=0)
        print('OK')
    finally:
        doc.Close(False)
        app.Quit()
finally:
    pythoncom.CoUninitialize()
"""
    params = json.dumps({
        "progid": progid,
        "input": os.path.abspath(input_path),
        "output": os.path.abspath(output_path),
    })

    result = subprocess.run(
        [sys.executable, "-c", script, params],
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0 or "OK" not in result.stdout:
        raise RuntimeError(f"COM export failed: {result.stderr or result.stdout}")


def _word_to_pdf_com(input_path: str, output_path: str) -> None:
    """Word→PDF via MS Word COM (deprecated, use _com_export)."""
    _com_export(input_path, output_path, "Word.Application")


def _word_to_pdf_libreoffice(input_path: str, output_path: str) -> None:
    """Word→PDF via LibreOffice headless CLI."""
    import subprocess
    out_dir = os.path.dirname(os.path.abspath(output_path)) or "."
    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf",
         "--outdir", out_dir, os.path.abspath(input_path)],
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
    # LibreOffice names the output after the input stem; rename if needed
    lo_output = os.path.join(out_dir, Path(input_path).stem + ".pdf")
    if os.path.abspath(lo_output) != os.path.abspath(output_path):
        import shutil
        shutil.move(lo_output, output_path)


def _word_to_pdf_pure(input_path: str, output_path: str) -> None:
    """Word→PDF: XML body-order traversal with element-identity matching."""

    import tempfile
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fpdf import FPDF

    doc = Document(input_path)
    pdf = FPDF()

    # ── Section formatting ──
    section = doc.sections[0]
    left_m = float(section.left_margin.inches * 25.4) if section.left_margin else 20
    right_m = float(section.right_margin.inches * 25.4) if section.right_margin else 20
    top_m = float(section.top_margin.inches * 25.4) if section.top_margin else 20
    pdf.set_left_margin(left_m)
    pdf.set_right_margin(right_m)
    pdf.set_auto_page_break(auto=True, margin=max(top_m, 15))
    pdf.add_page()

    # ── Font ──
    FONT = "Helvetica"
    for fp in ("C:/Windows/Fonts/simsun.ttc", "C:/Windows/Fonts/msyh.ttc"):
        try:
            pdf.add_font("CJK", "", fp)
            FONT = "CJK"
            break
        except Exception:
            pass

    # ── Element-identity maps (no index matching!) ──
    para_by_elem = {p._element: p for p in doc.paragraphs}
    table_by_elem = {t._element: t for t in doc.tables}

    # ── Image parts ──
    image_parts = {}
    for r in doc.part.rels.values():
        if "image" in r.reltype:
            image_parts[r.rId] = r.target_part

    # Namespaces
    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
    NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    NS_MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    NS_WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    NS_WPG = "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"

    def _pt_mm(pt_val):
        return pt_val * 0.3528

    def _all_text(elem):
        """Recursively extract all <w:t> text from an XML element."""
        parts = []
        for t_node in elem.iter(f"{{{NS_W}}}t"):
            if t_node.text:
                parts.append(t_node.text)
        return "".join(parts)

    def _render_images(elem):
        for blip in elem.iter(f"{{{NS_A}}}blip"):
            embed_id = blip.get(f"{{{NS_R}}}embed")
            if embed_id and embed_id in image_parts:
                img_part = image_parts[embed_id]
                try:
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
                        tf.write(img_part.blob)
                        tmp_path = tf.name
                    ext_cx = elem.find(f".//{{{NS_WP}}}extent")
                    max_w = pdf.w - left_m - right_m
                    img_w = max_w
                    if ext_cx is not None:
                        img_w = min(int(ext_cx.get("cx", "0")) / 914400 * 72, max_w)
                    pdf.image(tmp_path, x=(pdf.w - img_w) / 2, w=img_w)
                    pdf.ln(3)
                    os.unlink(tmp_path)
                except Exception:
                    pass

    def _render_para(para):
        """Render one paragraph with full formatting."""
        elem = para._element
        _render_images(elem)

        pf = para.paragraph_format
        space_before = _pt_mm(pf.space_before.pt) if pf.space_before else 0
        space_after = _pt_mm(pf.space_after.pt) if pf.space_after else 0
        line_spacing = pf.line_spacing
        first_indent = _pt_mm(pf.first_line_indent.pt) if pf.first_line_indent else 0

        if space_before > 0:
            pdf.ln(space_before)

        if not para.runs:
            pdf.ln(3)
            return

        # Collect run data
        run_data = []
        all_text = ""
        for run in para.runs:
            if not run.text:
                continue
            f = run.font
            sz = f.size.pt if f.size else None
            bold = f.bold or False
            italic = f.italic or False
            underline = f.underline or False
            color = None
            if f.color and f.color.rgb:
                try: color = (f.color.rgb[0], f.color.rgb[1], f.color.rgb[2])
                except Exception: pass
            run_data.append((run.text, sz, bold, italic, underline, color))
            all_text += run.text

        if not all_text.strip():
            pdf.ln(3)
            return

        # Style defaults
        sn = para.style.name.lower() if para.style else ""
        if sn.startswith("heading 1"): base_sz = 18
        elif sn.startswith("heading 2"): base_sz = 14
        elif "heading" in sn: base_sz = 12
        else:
            sf = para.style.font
            base_sz = sf.size.pt if sf and sf.size else 10

        # Alignment
        align = "L"
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER: align = "C"
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT: align = "R"
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY: align = "J"

        # Line height
        lh = base_sz * 0.55
        if line_spacing is not None:
            if isinstance(line_spacing, float):
                lh = base_sz * line_spacing * 0.55
            elif hasattr(line_spacing, 'pt'):
                lh = _pt_mm(line_spacing.pt)

        # Check uniform formatting
        first_sz = run_data[0][1]
        first_color = run_data[0][5]
        uniform = all(r[1] == first_sz and r[5] == first_color and not r[2] and not r[3] for r in run_data)

        if uniform and not first_indent:
            sz = first_sz or base_sz
            pdf.set_font(FONT, "", sz)
            if first_color: pdf.set_text_color(*first_color)
            else: pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(0, lh, all_text, align=align)
        else:
            # Mixed formatting with write()
            pdf.set_font(FONT, "", base_sz)
            if first_indent > 0:
                pdf.cell(first_indent, lh, "")
            for text, sz, bold, italic, underline, color in run_data:
                sz = sz or base_sz
                style = ""
                if bold and FONT == "Helvetica": style += "B"
                if italic and FONT == "Helvetica": style += "I"
                if underline and FONT == "Helvetica": style += "U"
                pdf.set_font(FONT, style, sz)
                pdf.set_text_color(*(color or (0, 0, 0)))
                pdf.write(lh, text)
            pdf.ln(lh + 0.5)

        if space_after > 0:
            pdf.ln(space_after)

    def _render_table(table):
        pdf.set_font(FONT, "", 8)
        pdf.set_text_color(0, 0, 0)
        cols = len(table.columns)
        col_w = (pdf.w - left_m - right_m) / max(cols, 1)
        for row in table.rows:
            for cell in row.cells:
                pdf.cell(col_w, 7, cell.text[:80], border=1)
            pdf.ln()
        pdf.ln(4)

    def _render_plain_text(text):
        """Fallback: render extracted text as a simple paragraph."""
        text = text.strip()
        if not text:
            return
        pdf.set_font(FONT, "", 10)
        pdf.set_text_color(0, 0, 0)
        pdf.multi_cell(0, 5.5, text)
        pdf.ln(1)

    # ── Process body children recursively ──
    def _process_children(parent):
        for child in parent:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                para = para_by_elem.get(child)
                if para is not None:
                    _render_para(para)
                else:
                    # Paragraph not in doc.paragraphs? Extract text anyway.
                    _render_plain_text(_all_text(child))

            elif tag == "tbl":
                table = table_by_elem.get(child)
                if table is not None:
                    _render_table(table)
                else:
                    _render_plain_text(_all_text(child))

            elif tag == "sdt":
                sdt_content = child.find(f"{{{NS_W}}}sdtContent")
                if sdt_content is not None:
                    _process_children(sdt_content)

            elif tag == "AlternateContent":
                # Handle mc:AlternateContent — pick the Choice or Fallback
                choice = child.find(f"{{{NS_MC}}}Choice")
                if choice is not None:
                    _process_children(choice)
                else:
                    fallback = child.find(f"{{{NS_MC}}}Fallback")
                    if fallback is not None:
                        _process_children(fallback)

            elif tag == "std":
                # <w:std> elements (custom XML) — recurse
                _process_children(child)

            else:
                # ── Text boxes (txbxContent) ──
                txbx_contents = child.iter(f"{{{NS_WPS}}}txbxContent")
                for txbx in txbx_contents:
                    # Process paragraphs inside the text box
                    for p_elem in txbx.iter(f"{{{NS_W}}}p"):
                        para = para_by_elem.get(p_elem)
                        if para:
                            _render_para(para)
                        else:
                            txt = _all_text(p_elem)
                            if txt.strip():
                                _render_plain_text(txt)

                # ── Fallback: extract any remaining text ──
                txt = _all_text(child)
                if txt.strip():
                    _render_images(child)
                    _render_plain_text(txt)

    # ── Safety font ──
    pdf.set_font(FONT, "", 10)
    pdf.set_text_color(0, 0, 0)

    _process_children(doc.element.body)

    pdf.output(output_path)


def pdf_to_word(input_path: str, output_path: str) -> None:
    """Convert a .pdf file to .docx using pdf2docx."""
    try:
        from pdf2docx import Converter
    except ImportError:
        print("[error] pdf2docx is required for PDF->Word conversion.")
        print("        Install it with: pip install pdf2docx")
        sys.exit(1)

    cv = Converter(input_path)
    cv.convert(output_path)       # all pages by default
    cv.close()


# ── Image conversion ─────────────────────────────────────────────────

IMAGE_FORMATS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tiff", ".tif"}


def image_convert(input_path: str, output_path: str) -> None:
    """Convert between image formats using Pillow."""
    from PIL import Image

    img = Image.open(input_path)
    out_ext = Path(output_path).suffix.lower()

    # JPEG doesn't support alpha channel
    if out_ext in (".jpg", ".jpeg") and img.mode in ("RGBA", "LA", "P"):
        img = img.convert("RGB")
    # GIF needs special handling
    elif out_ext == ".gif" and getattr(img, "is_animated", False):
        img.save(output_path, save_all=True, loop=0)
        return

    img.save(output_path)
    print(f"Image converted: {Path(input_path).suffix} → {out_ext}  ({img.size[0]}x{img.size[1]})")

def auto_output_path(input_path: str) -> str:
    """Derive output path from input path by switching extension."""
    p = Path(input_path)
    ext = p.suffix.lower()
    if ext == ".docx":
        return str(p.with_suffix(".pdf"))
    elif ext == ".pdf":
        return str(p.with_suffix(".docx"))
    elif ext in IMAGE_FORMATS:
        return str(p.with_suffix(".png" if ext != ".png" else ".jpg"))
    else:
        raise ValueError(f"Unsupported input type: {ext}. Expected .docx, .pdf, or image format")


def convert(input_path: str, output_path: Optional[str] = None) -> str:
    """Main conversion entry point. Returns the output path."""
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")

    if output_path is None:
        output_path = auto_output_path(input_path)
    elif os.path.isdir(output_path):
        name = Path(input_path).stem
        ext = Path(input_path).suffix.lower()
        if ext in IMAGE_FORMATS:
            out_ext = ".png" if ext != ".png" else ".jpg"
        elif ext == ".docx":
            out_ext = ".pdf"
        else:
            out_ext = ".docx"
        output_path = os.path.join(output_path, name + out_ext)

    in_ext = Path(input_path).suffix.lower()
    out_ext = Path(output_path).suffix.lower()

    if in_ext == out_ext:
        raise ValueError("Input and output have the same extension. Nothing to convert.")

    print(f"Converting: {input_path}  →  {output_path}")

    # Route to correct converter
    if in_ext == ".docx" and out_ext == ".pdf":
        word_to_pdf(input_path, output_path)
    elif in_ext == ".pdf" and out_ext == ".docx":
        pdf_to_word(input_path, output_path)
    elif in_ext in IMAGE_FORMATS and out_ext in IMAGE_FORMATS:
        image_convert(input_path, output_path)
    else:
        raise ValueError(
            f"Unsupported conversion: {in_ext} → {out_ext}. "
            f"Supports: docx↔pdf, and image format interchange"
        )

    print(f"Done: {output_path}")
    return output_path


# ── GUI mode ──────────────────────────────────────────────────────────

def run_gui():
    """Simple Tkinter GUI with file picker and drag-drop support."""
    root = tk.Tk()
    root.title("文档 & 图片互转工具")
    root.geometry("520x320")
    root.resizable(False, False)

    # Center the window
    root.update_idletasks()
    w, h = root.winfo_width(), root.winfo_height()
    sw, sh = root.winfo_screenwidth(), root.winfo_screenheight()
    root.geometry(f"+{(sw-w)//2}+{(sh-h)//2}")

    # Style
    style = ttk.Style()
    style.theme_use("clam")

    # State
    selected_file = tk.StringVar()
    direction = tk.StringVar(value="auto")

    def pick_file():
        path = filedialog.askopenfilename(
            title="选择文件",
            filetypes=[
                ("所有支持格式", "*.docx;*.pdf;*.png;*.jpg;*.jpeg;*.webp;*.bmp;*.gif;*.tiff;*.tif"),
                ("Word & PDF 文档", "*.docx;*.pdf"),
                ("图片文件", "*.png;*.jpg;*.jpeg;*.webp;*.bmp;*.gif;*.tiff;*.tif"),
                ("Word 文档", "*.docx"),
                ("PDF 文件", "*.pdf"),
            ],
        )
        if path:
            selected_file.set(path)
            update_preview()

    def update_preview(*_args):
        path = selected_file.get()
        if not path:
            preview_var.set("未选择文件")
            return
        p = Path(path)
        ext = p.suffix.lower()
        if ext == ".docx":
            out = p.with_suffix(".pdf")
            preview_var.set(f"输入: {p.name}\n输出: {out.name}\n方向: Word → PDF")
        elif ext == ".pdf":
            out = p.with_suffix(".docx")
            preview_var.set(f"输入: {p.name}\n输出: {out.name}\n方向: PDF → Word")
        elif ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tiff", ".tif"}:
            out_ext = ".png" if ext != ".png" else ".jpg"
            out = p.with_suffix(out_ext)
            preview_var.set(f"输入: {p.name}\n输出: {out.name}\n方向: {ext} → {out_ext}")
        else:
            preview_var.set(f"不支持的文件类型: {p.suffix}")

    def do_convert():
        path = selected_file.get()
        if not path:
            messagebox.showwarning("提示", "请先选择一个文件")
            return
        try:
            out = convert(path)
            messagebox.showinfo("完成", f"转换成功!\n\n输出文件:\n{out}")
        except Exception as e:
            messagebox.showerror("错误", str(e))

    # Layout
    frame = ttk.Frame(root, padding=24)
    frame.pack(fill="both", expand=True)

    ttk.Label(frame, text="文档 & 图片互转工具", font=("Segoe UI", 16, "bold")).pack(pady=(0, 16))

    ttk.Button(frame, text="📂 选择文件", command=pick_file).pack(pady=4)

    preview_var = tk.StringVar(value="未选择文件")
    preview_lbl = ttk.Label(
        frame, textvariable=preview_var, font=("Consolas", 10),
        justify="left", anchor="w",
        background="#f0f0f0", padding=12,
    )
    preview_lbl.pack(fill="x", pady=(12, 20))

    ttk.Button(frame, text="🔀 开始转换", command=do_convert).pack()

    root.mainloop()


def main():
    parser = argparse.ArgumentParser(
        description="文档 & 图片互转工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    py doc_pdf_converter.py report.docx            # → report.pdf
    py doc_pdf_converter.py report.pdf             # → report.docx
    py doc_pdf_converter.py photo.jpg              # → photo.png
    py doc_pdf_converter.py input.docx -o out.pdf
    py doc_pdf_converter.py --gui                  # Launch GUI
        """,
    )
    parser.add_argument("input", nargs="?", help="Input file (.docx, .pdf, or image)")
    parser.add_argument("-o", "--output", default=None, help="Output file path")
    parser.add_argument("--gui", action="store_true", help="Launch graphical interface")

    args = parser.parse_args()

    if args.gui:
        run_gui()
        return

    if not args.input:
        parser.print_help()
        print("\n[tip] Use --gui for the graphical interface, or drag a file onto this script.")
        sys.exit(1)

    try:
        convert(args.input, args.output)
    except Exception as e:
        print(f"[error] {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
